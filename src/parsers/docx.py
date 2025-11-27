"""
Word document (.docx) parser.

Uses python-docx to extract text from Word documents.
Extracts text from paragraphs and tables.

Note:
    - Only works with .docx files (Office 2007+)
    - Old .doc files are not supported (would need different library)
    - Headers, footers, and comments are not extracted
"""

from pathlib import Path
from typing import Union

from docx import Document


def parse_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from a Word document (.docx).

    Args:
        file_path: Path to the .docx file

    Returns:
        Extracted text from paragraphs and tables

    Example:
        >>> text = parse_docx("resume.docx")
        >>> print(text[:100])
        "John Doe\\nSoftware Engineer\\n..."
    """
    file_path = Path(file_path)
    doc = Document(file_path)

    text_parts = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


# Test the parser directly
if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        text = parse_docx(docx_path)
        print(f"Extracted {len(text)} characters from {docx_path}")
        print("\n--- First 500 characters ---")
        print(text[:500])
    else:
        print("Usage: python -m src.parsers.docx <path_to_docx>")
